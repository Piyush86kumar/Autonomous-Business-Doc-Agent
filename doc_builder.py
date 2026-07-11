"""Deterministic .docx assembly from drafted sections."""

from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from schemas import AgentState

OUTPUT_DIR = Path(__file__).parent / "outputs"
ACCENT_COLOR = RGBColor(0x1F, 0x4E, 0x79)

_LINE_KV_PATTERN = re.compile(r"^\s*[-*]?\s*([^:\n]{1,60}):\s*(.+)$")


def _maybe_build_table(document: Document, content: str) -> bool:
    """Render "Key: Value" lines as a two-column table. Returns True if table was rendered."""
    lines = [ln for ln in content.splitlines() if ln.strip()]
    kv_pairs = []
    for line in lines:
        match = _LINE_KV_PATTERN.match(line)
        if not match:
            return False
        kv_pairs.append((match.group(1).strip(), match.group(2).strip()))

    if len(kv_pairs) < 2:
        return False

    table = document.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text = "Item", "Detail"
    for key, value in kv_pairs:
        row = table.add_row().cells
        row[0].text, row[1].text = key, value
    document.add_paragraph()  # spacing after table
    return True


def build_document(state: AgentState) -> str:
    assert state.outline is not None, "Doc builder requires a completed outline + sections"

    document = Document()

    title = document.add_heading(state.outline.title, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        f"{state.outline.doc_type.replace('_', ' ').title()}  |  "
        f"Generated {datetime.now().strftime('%B %d, %Y')}"
    )
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    document.add_paragraph()

    for planned in state.outline.sections:
        section = state.sections.get(planned.section_name)
        if section is None:
            continue

        document.add_heading(section.section_name, level=1)

        rendered_as_table = False
        if section.needs_table:
            rendered_as_table = _maybe_build_table(document, section.content)

        if not rendered_as_table:
            for para_text in section.content.split("\n\n"):
                para_text = para_text.strip()
                if not para_text:
                    continue
                document.add_paragraph(para_text)

    if state.assumptions_made:
        document.add_heading("Assumptions Made", level=1)
        note = document.add_paragraph()
        note.add_run(
            "The original request did not fully specify the following — "
            "reasonable assumptions were made and are disclosed here rather "
            "than silently embedded in the document above:"
        ).italic = True
        for assumption in state.assumptions_made:
            document.add_paragraph(assumption, style="List Bullet")

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    safe_type = re.sub(r"[^a-z0-9_]+", "", state.outline.doc_type.lower().replace(" ", "_"))
    filename = f"{safe_type or 'document'}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = OUTPUT_DIR / filename

    document.save(str(filepath))
    return str(filepath)
